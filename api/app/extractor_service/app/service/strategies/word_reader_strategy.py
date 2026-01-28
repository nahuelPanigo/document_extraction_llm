from docx import Document
from app.service.strategies.reader_strategy import ReaderStrategy
import tempfile
import os
from PIL import Image
import numpy as np
import io


class DocxReader(ReaderStrategy):
    def __new__(cls):
        if not hasattr(cls, 'instance'):
            cls.instance = super(DocxReader, cls).__new__(cls)
        return cls.instance

    def get_correct_tag(self, font_size: int, sizes: dict) -> str:
        if font_size >= sizes['h1']:
            return "h1"
        if font_size >= sizes['h2']:
            return "h2"
        return "p"

    def _extract_ocr_from_docx(self, docx_path: str, ocr_reader, wrap_tags=True):
        """
        Extract text from images in DOCX file using OCR

        Args:
            docx_path: Path to DOCX file
            ocr_reader: EasyOCR reader instance
            wrap_tags: Whether to wrap OCR text in <img> tags

        Returns:
            List of OCR text strings
        """
        doc = Document(docx_path)
        ocr_texts = []
        processed_sizes = set()

        # Extract images from document relationships
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    img = Image.open(io.BytesIO(image_data))

                    img_width, img_height = img.size
                    size_key = (img_width, img_height)

                    # Skip if already processed (duplicate detection)
                    if size_key in processed_sizes:
                        print(f"🔄 Skipping duplicate image {size_key}")
                        continue

                    # Skip very small images (likely icons or decorations)
                    if img_width < 100 or img_height < 100:
                        print(f"⏭️ Skipping small image {size_key}")
                        continue

                    # Process image with OCR
                    print(f"🔍 Processing image {size_key} with EasyOCR...")
                    img_array = np.array(img.convert('RGB'))
                    results = ocr_reader.readtext(img_array)
                    text = ' '.join([result[1] for result in results])

                    if text.strip():
                        if wrap_tags:
                            ocr_texts.append(f"<img>{text.strip()}</img>")
                        else:
                            ocr_texts.append(text.strip())
                        print(f"✅ Extracted: {text[:50]}...")

                    # Mark this size as processed
                    processed_sizes.add(size_key)

                except Exception as e:
                    print(f"❌ Error extracting image: {e}")
                    continue

        return ocr_texts

    def get_fontsizes(self, docx_path: str) -> dict:
        doc = Document(docx_path)
        sizes = []
        for para in doc.paragraphs[:100]:  # limitar a los primeros 100 párrafos
            for run in para.runs:
                if run.font.size:
                    pt_size = run.font.size.pt  # convertir a puntos
                    if pt_size < 40:
                        sizes.append(round(pt_size))

        sizes = sorted(set(sizes))
        if not sizes:
            return {'h1': 16, 'h2': 14, 'p': 9}  # defaults si no hay tamaños

        h1_size = sizes[-1]
        n = len(sizes)
        h2_size = sizes[int(n * 0.75)] if n > 1 else sizes[0]
        paragraph_sizes = [size for size in sizes if size < h2_size]
        return {
            'h1': h1_size,
            'h2': h2_size,
            'p': min(paragraph_sizes) if paragraph_sizes else sizes[0]
        }


    def extract_text_with_xml_tags(self, docx_path:str, ocr: bool = False) -> str:
        """
        Extract text with XML tags and optionally OCR from images

        Args:
            docx_path: Path to DOCX file
            ocr: Boolean flag to enable/disable OCR processing
        """
        sizes_dict = self.get_fontsizes(docx_path)
        doc = Document(docx_path)

        # Initialize OCR if enabled
        ocr_reader = None
        if ocr:
            try:
                import easyocr
                print("🔧 Initializing EasyOCR for image text extraction...")
                ocr_reader = easyocr.Reader(['en', 'es'])
            except ImportError:
                print("❌ EasyOCR not available. Install with: pip install easyocr")
                ocr = False

        text_with_tags = ""
        current_tag = "p"
        current_text = []

        for para in doc.paragraphs:
            para_text = ""
            max_font_size = 0
            for run in para.runs:
                if run.font.size:
                    font_size = round(run.font.size.pt)
                    max_font_size = max(max_font_size, font_size)
                para_text += run.text

            if para_text.strip():
                tag = self.get_correct_tag(max_font_size, sizes_dict)
                if current_tag != tag:
                    if current_text:
                        text_with_tags += f"<{current_tag}>" + " ".join(current_text) + f"</{current_tag}>"
                    current_text = [para_text.strip()]
                    current_tag = tag
                else:
                    current_text.append(para_text.strip())

                if len(text_with_tags.split(" ")) > 2000:
                    break

        if current_text:
            text_with_tags += f"<{current_tag}>" + " ".join(current_text) + f"</{current_tag}>"

        # Extract OCR text from images if enabled
        if ocr and ocr_reader:
            ocr_texts = self._extract_ocr_from_docx(docx_path, ocr_reader, wrap_tags=True)
            if ocr_texts:
                text_with_tags += ''.join(ocr_texts)

        return text_with_tags
    
    def extract_text(self, docx_path: str, ocr: bool = False) -> str:
        """
        Extract plain text from DOCX and optionally OCR from images

        Args:
            docx_path: Path to DOCX file
            ocr: Boolean flag to enable/disable OCR processing
        """
        doc = Document(docx_path)
        lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

        # Extract OCR text from images if enabled
        if ocr:
            try:
                import easyocr
                print("🔧 Initializing EasyOCR for image text extraction...")
                ocr_reader = easyocr.Reader(['en', 'es'])

                ocr_texts = self._extract_ocr_from_docx(docx_path, ocr_reader, wrap_tags=False)
                if ocr_texts:
                    lines.extend(ocr_texts)
            except ImportError:
                print("❌ EasyOCR not available. Install with: pip install easyocr")

        return "\n".join(lines)
